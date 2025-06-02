import csv
import json
import shutil
import subprocess
import tarfile
import zipfile
from pathlib import Path

import cairosvg
import fitz  # PyMuPDF
import svgwrite
from PIL import Image
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


class Base:
    conversion_formats = None

    def __init__(self):
        if self.conversion_formats is None:
            raise NotImplementedError("Subclasses must define `conversion_formats` and `class_func`")

    def convert(self, source_format, target_format, input_path, output_path):
        method_name = f'{source_format}_to_{target_format}'
        method = getattr(self.__class__, method_name, None)
        return method(input_path, output_path)

    def __init_subclass__(cls, **kwargs):
        if getattr(cls, 'conversion_formats') is None:
            raise AttributeError('Любой интерфейс должен определить словарь conversion_formats')

        for source, targets in cls.conversion_formats.items():
            for target in targets:
                method_name = f'{source}_to_{target}'
                if getattr(cls, method_name) is None:
                    raise AttributeError('Не задан метод для конвертации {source} => {targets}')


class Image(Base):
    conversion_formats = {
        'jpg': ['webp', 'svg'],
        'png': ['webp', 'svg'],
        'webp': ['png', 'jpg'],
        'svg': ['png']
    }

    @staticmethod
    def jpg_to_webp(input_path: Path, output_path: Path):
        with Image.open(input_path) as img:
            img.save(output_path, "WEBP")

    @staticmethod
    def png_to_webp(input_path: Path, output_path: Path):
        with Image.open(input_path) as img:
            img.save(output_path, "WEBP")

    @staticmethod
    def webp_to_png(input_path: Path, output_path: Path):
        with Image.open(input_path) as img:
            img = img.convert("RGB")
            img.save(output_path, "PNG")

    @staticmethod
    def webp_to_jpg(input_path: Path, output_path: Path):
        with Image.open(input_path) as img:
            img = img.convert("RGB")
            img.save(output_path, "JPEG")

    @staticmethod
    def svg_to_png(input_path: Path, output_path: Path):
        cairosvg.svg2png(url=str(input_path), write_to=str(output_path))

    @staticmethod
    def raster_to_svg(input_path: Path, output_path: Path, threshold: int = 128):
        img = Image.open(input_path).convert("L")  # Grayscale
        width, height = img.size
        pixels = img.load()

        dwg = svgwrite.Drawing(str(output_path), profile='tiny', size=(width, height))

        for y in range(height):
            for x in range(width):
                if pixels[x, y] < threshold:  # Чёрная точка
                    dwg.add(dwg.rect(insert=(x, y), size=(1, 1), fill='black'))

        dwg.save()

    @staticmethod
    def jpg_to_svg(input_path: Path, output_path: Path):
        Image.raster_to_svg(input_path, output_path)

    @staticmethod
    def png_to_svg(input_path: Path, output_path: Path):
        Image.raster_to_svg(input_path, output_path)


class Text(Base):
    conversion_formats = {
        'txt': ['csv', 'json', 'pdf', 'docx'],
        'csv': ['txt'],
        'json': ['txt'],
        'pdf': ['txt'],
        'docx': ['txt']
    }

    @staticmethod
    def txt_to_csv(input_path: Path, output_path: Path):
        content = input_path.read_text(encoding="utf-8")
        lines = content.strip().splitlines()
        rows = [line.split() for line in lines]
        with output_path.open("w", newline='', encoding="utf-8") as csvfile:
            writer = csv.writer(csvfile)
            writer.writerows(rows)

    @staticmethod
    def txt_to_json(input_path: Path, output_path: Path):
        content = input_path.read_text(encoding="utf-8")
        lines = content.strip().splitlines()
        data = {"lines": lines}
        with output_path.open("w", encoding="utf-8") as jsonfile:
            json.dump(data, jsonfile, ensure_ascii=False, indent=2)

    @staticmethod
    def csv_to_txt(input_path: Path, output_path: Path):
        with input_path.open("r", encoding="utf-8") as csvfile:
            reader = csv.reader(csvfile)
            with output_path.open("w", encoding="utf-8") as txtfile:
                for row in reader:
                    txtfile.write(" ".join(row) + "\n")

    @staticmethod
    def json_to_txt(input_path: Path, output_path: Path):
        content = input_path.read_text(encoding="utf-8")
        data = json.loads(content)
        with output_path.open("w", encoding="utf-8") as txtfile:
            if isinstance(data, dict):
                for key, value in data.items():
                    txtfile.write(f"{key}: {value}\n")
            elif isinstance(data, list):
                for item in data:
                    txtfile.write(f"{item}\n")

    @staticmethod
    def txt_to_pdf(input_path: Path, output_path: Path) -> None:
        with input_path.open('r', encoding='utf-8') as file:
            lines = file.readlines()

        c = canvas.Canvas(str(output_path), pagesize=A4)
        width, height = A4
        y = height - 40

        for line in lines:
            if y < 40:
                c.showPage()
                y = height - 40
            c.drawString(40, y, line.strip())
            y -= 15

        c.save()

    @staticmethod
    def pdf_to_txt(input_path: Path, output_path: Path):
        text = ""
        with fitz.open(input_path) as pdf:
            for page in pdf:
                text += page.get_text()
        output_path.write_text(text, encoding="utf-8")

    @staticmethod
    def docx_to_txt(input_path: Path, output_path: Path):
        document = Document(input_path)
        with output_path.open("w", encoding="utf-8") as file:
            for para in document.paragraphs:
                file.write(para.text + "\n")

    @staticmethod
    def txt_to_docx(input_path: Path, output_path: Path):
        document = Document()
        with input_path.open("r", encoding="utf-8") as file:
            for line in file:
                document.add_paragraph(line.strip())
        document.save(output_path)


class Video(Base):
    conversion_formats = {
        'mp4': ['avi', 'webm', 'mov'],
        'avi': ['mp4'],
        'webm': ['mp4'],
        'mov': ['mp4']
    }

    @classmethod
    def run(cls, params: list, input_path: Path, output_path: Path):
        command = ["ffmpeg", "-i", str(input_path)] + params + [str(output_path)]
        subprocess.run(command, check=True)

    @classmethod
    def mp4_to_avi(cls, input_path, output_path):
        params = ['-c:v', 'mpeg4']
        return Video.run(params, input_path, output_path)

    @classmethod
    def mp4_to_webm(cls, input_path, output_path):
        params = ['-c:v', 'libvpx', '-b:v', '1M', '-c:a', 'libvorbis']
        return Video.run(params, input_path, output_path)

    @classmethod
    def mp4_to_mov(cls, input_path, output_path):
        params = ['-c:v', 'libx264', '-c:a', 'aac']
        return Video.run(params, input_path, output_path)

    @classmethod
    def avi_to_mp4(cls, input_path, output_path):
        params = ['-c:v', 'libx264', '-c:a', 'aac']
        return Video.run(params, input_path, output_path)

    @classmethod
    def webm_to_mp4(cls, input_path, output_path):
        params = ['-c:v', 'libx264', '-c:a', 'aac']
        return Video.run(params, input_path, output_path)

    @classmethod
    def mov_to_mp4(cls, input_path, output_path):
        params = ['-c:v', 'libx264', '-c:a', 'aac']
        return Video.run(params, input_path, output_path)


class DecompressingError(Exception):
    pass


class Archive(Base):
    conversion_formats = {
        'folder': ['zip', 'tar_gz'],  # tag_gz вполне уместная запись
        'zip': ['folder'],
        'tar_gz': ['folder']
    }

    @staticmethod
    def normalize_format(format: str):
        return format.replace('tar.gz', 'tar_gz')

    def convert(self, source_format: str, target_format: str, input_path: Path, output_path: Path):
        source_format, target_format = self.normalize_format(source_format), self.normalize_format(target_format)
        return super().convert(source_format, target_format, input_path, output_path)

    @staticmethod
    def folder_to_zip(input_path: Path, output_path: Path):
        shutil.make_archive(str(output_path.with_suffix('')), 'zip', root_dir=input_path)

    @staticmethod
    def folder_to_tar_gz(input_path: Path, output_path: Path):
        with tarfile.open(output_path, "w:gz") as tar:
            tar.add(input_path, arcname=".")

    @staticmethod
    def zip_to_folder(input_path: Path, output_path: Path):
        with zipfile.ZipFile(input_path, 'r') as zip_ref:
            zip_ref.extractall(output_path)

    @staticmethod
    def tar_gz_to_folder(input_path: Path, output_path: Path):
        with tarfile.open(input_path, 'r:gz') as tar:
            tar.extractall(path=output_path)


class Selector:  # Selector
    INTERFACES = [Text, Video, Image, Archive]

    def __init__(self, source_format: str, target_format: str, input_path: Path, output_path: Path):
        self.source_format = source_format
        self.target_format = target_format
        self.input_path = input_path
        self.output_path = output_path
        self.interface = self.resolve_interface()

    def resolve_interface(self):
        method_name = f'{self.source_format}_to_{self.target_format}'
        for iface_cls in self.INTERFACES:
            if (
                    self.source_format in iface_cls.conversion_formats
                    and self.target_format in iface_cls.conversion_formats[self.source_format]
                    and hasattr(iface_cls, method_name)
            ):
                return iface_cls()
        raise ValueError(
            f"Формат '{self.source_format}' → '{self.target_format}' не поддерживается ни одним конвертором.")

    def run(self) -> bool:
        return self.interface.convert(self.source_format, self.target_format, self.input_path, self.output_path)

    @classmethod
    def get_all_supported_pairs(cls):
        all_pairs = []
        for iface in cls.INTERFACES:
            _conversion_formats = iface.conversion_formats
            format_pairs = [(key, value) for key, values in _conversion_formats.items() for value in values]
            all_pairs.extend(format_pairs)
        return all_pairs
