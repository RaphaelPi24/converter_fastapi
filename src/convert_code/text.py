import csv
import json
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def txt_to_csv(input_path: Path, output_path: Path):
    content = input_path.read_text(encoding="utf-8")
    lines = content.strip().splitlines()
    rows = [line.split() for line in lines]
    with output_path.open("w", newline='', encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerows(rows)


def txt_to_json(input_path: Path, output_path: Path):
    content = input_path.read_text(encoding="utf-8")
    lines = content.strip().splitlines()
    data = {"lines": lines}
    with output_path.open("w", encoding="utf-8") as jsonfile:
        json.dump(data, jsonfile, ensure_ascii=False, indent=2)


def csv_to_txt(input_path: Path, output_path: Path):
    with input_path.open("r", encoding="utf-8") as csvfile:
        reader = csv.reader(csvfile)
        with output_path.open("w", encoding="utf-8") as txtfile:
            for row in reader:
                txtfile.write(" ".join(row) + "\n")


def json_to_txt(input_path: Path, output_path: Path):
    content = input_path.read_text(encoding="utf-8")
    data = json.loads(content)
    with output_path.open("w", encoding="utf-8") as txtfile:
        if isinstance(data, dict):
            for key, value in data.items():
                txtfile.write(f"{key}: {value}\n")
        elif isinstance(data, list):
            for item in data:
                txtfile.write(f"{item}\n")


def txt_to_pdf(input_path: Path, output_path: Path) -> None:
    with input_path.open('r', encoding='utf-8') as file:
        lines = file.readlines()

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    y = height - 40

    for line in lines:
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(40, y, line.strip())
        y -= 15

    c.save()


def pdf_to_txt(input_path: Path, output_path: Path):
    text = ""
    with fitz.open(input_path) as pdf:
        for page in pdf:
            text += page.get_text()
    output_path.write_text(text, encoding="utf-8")


def docx_to_txt(input_path: Path, output_path: Path):
    document = Document(input_path)
    with output_path.open("w", encoding="utf-8") as file:
        for para in document.paragraphs:
            file.write(para.text + "\n")


def txt_to_docx(input_path: Path, output_path: Path):
    document = Document()
    with input_path.open("r", encoding="utf-8") as file:
        for line in file:
            document.add_paragraph(line.strip())
    document.save(output_path)
